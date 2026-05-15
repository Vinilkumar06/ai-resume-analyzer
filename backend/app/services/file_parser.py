import io
import PyPDF2
import pdfplumber
from docx import Document
from loguru import logger
from fastapi import HTTPException


async def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber with PyPDF2 fallback."""
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            return text.strip()
    except Exception as e:
        logger.warning(f"pdfplumber failed: {e}, trying PyPDF2...")

    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
    except Exception as e:
        logger.error(f"PyPDF2 also failed: {e}")
        raise HTTPException(status_code=400, detail="Could not extract text from PDF")

    if not text.strip():
        raise HTTPException(status_code=400, detail="PDF appears to be empty or image-based")
    return text.strip()


async def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX files."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        text = "\n".join(paragraphs)
        if not text.strip():
            raise HTTPException(status_code=400, detail="DOCX appears to be empty")
        return text.strip()
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise HTTPException(status_code=400, detail="Could not extract text from DOCX")


async def extract_text(file_bytes: bytes, content_type: str, filename: str) -> str:
    """Route to correct extractor based on file type."""
    fn_lower = filename.lower()
    if fn_lower.endswith(".pdf") or content_type == "application/pdf":
        return await extract_text_from_pdf(file_bytes)
    elif fn_lower.endswith(".docx") or "word" in content_type:
        return await extract_text_from_docx(file_bytes)
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type. Use PDF or DOCX.")
